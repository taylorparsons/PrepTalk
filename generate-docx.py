#!/usr/bin/env python3
"""
PrepTalk UI/UX Strategic Proposal - Executive Document

Vision: Transform the entire PrepTalk experience from interview simulator to confidence builder
Audience: Business partners, UI/UX engineers, designers
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DIAGRAMS_PATH = 'docs/diagrams'
WIREFRAMES_PATH = 'docs/diagrams/wireframes'
STORYBOARD_PATH = 'docs/diagrams/storyboard'
OUTPUT_PATH = 'docs/PrepTalk_UI_UX_Proposal.docx'


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
    doc.add_paragraph()
    return table


def add_image(doc, path, caption, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.italic = True
            run.font.size = Pt(10)
        doc.add_paragraph()
    else:
        doc.add_paragraph(f'[Image not found: {path}]')


def title_page(doc):
    title = doc.add_heading('PrepTalk', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    subtitle = doc.add_paragraph('UI/UX Strategic Proposal')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(18)

    doc.add_paragraph()

    tagline = doc.add_paragraph('From Interview Simulator to Confidence Builder')
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in tagline.runs:
        run.italic = True
        run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    doc.add_paragraph('Prepared for: Business Partners, UI/UX Engineers & Designers')
    doc.add_paragraph('Date: February 3, 2026')

    doc.add_page_break()


def executive_summary(doc):
    add_heading(doc, 'Executive Summary')

    add_heading(doc, 'What is PrepTalk?', level=2)
    doc.add_paragraph(
        'PrepTalk is a voice-first interview practice platform. Users upload their resume '
        'and job description, receive tailored interview questions, practice answering via voice, '
        'get personalized coaching, and export study guides. All coaching is grounded in the '
        'user\'s actual experience—PrepTalk never fabricates.'
    )

    add_heading(doc, 'Who is it for?', level=2)
    doc.add_paragraph(
        'Anxious job seekers who have the experience but freeze under pressure. '
        'They don\'t need more interview questions—they need help discovering which of their '
        'experiences to draw from when answering.'
    )

    add_heading(doc, 'The Strategic Opportunity', level=2)
    doc.add_paragraph(
        'PrepTalk currently operates as an interview simulator: it tests users and offers '
        'coaching only after they struggle. This proposal recommends repositioning the entire '
        'experience as a confidence builder: teaching users to find their story first, '
        'then letting them practice with that foundation.'
    )

    add_heading(doc, 'The Decision', level=2)
    p = doc.add_paragraph()
    p.add_run(
        'Should PrepTalk lead with teaching or testing?'
    ).italic = True

    doc.add_page_break()


def the_user(doc):
    add_heading(doc, 'The User')

    quote = doc.add_paragraph()
    quote.add_run(
        '"I have the experience, but when they ask me a question, I freeze. '
        'I don\'t know what story to tell."'
    ).italic = True
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    add_heading(doc, 'The Problem Isn\'t Knowledge', level=2)
    doc.add_paragraph(
        'Our target user isn\'t unprepared. They have relevant experience. '
        'The problem is an ACCESS gap: their experiences exist in memory but aren\'t '
        'connected to the question being asked. Interview anxiety blocks retrieval.'
    )

    add_heading(doc, 'What They Need', level=2)
    add_table(doc,
        ['They Don\'t Need', 'They Need'],
        [
            ['More practice questions', 'Help connecting their experience to questions'],
            ['To be tested', 'To be coached'],
            ['Judgment on performance', 'Guidance on narrative'],
            ['Generic advice', 'Their own story reflected back'],
        ])

    add_heading(doc, 'The Opportunity', level=2)
    doc.add_paragraph(
        'PrepTalk already has the capability to show users exactly which of their experiences '
        'answers each question. The opportunity is to make this the core experience, '
        'not an afterthought.'
    )

    doc.add_page_break()


def design_principles(doc):
    add_heading(doc, 'Design Principles')

    doc.add_paragraph(
        'These principles should guide every UI/UX decision in PrepTalk:'
    )

    principles = [
        ('Coach, Not Judge',
         'Every interaction should feel like working with a supportive coach, '
         'not taking a test. Feedback is guidance, not grades.'),

        ('Teach Before Test',
         'Help users discover their narrative before asking them to perform. '
         'Show them the path, then let them walk it.'),

        ('Their Story, Not Ours',
         'All coaching draws from the user\'s actual resume. '
         'Never fabricate. Their experience is enough.'),

        ('Reduce Anxiety, Don\'t Add It',
         'Every UI element should reduce cognitive load. '
         'Simplify choices. Provide clear next steps. Celebrate progress.'),

        ('Voice-First, Not Voice-Only',
         'Optimize for speaking practice, but support users who need to read, '
         'review, or take notes.'),
    ]

    for title, description in principles:
        p = doc.add_paragraph()
        p.add_run(f'{title}: ').bold = True
        p.add_run(description)

    doc.add_page_break()


def journey_overview(doc):
    add_heading(doc, 'The User Journey')

    doc.add_paragraph(
        'PrepTalk has five phases. This section analyzes each phase and recommends '
        'how to shift from testing to teaching throughout the experience.'
    )

    add_table(doc,
        ['Phase', 'Current Framing', 'Proposed Framing'],
        [
            ['1. Arrival', 'Another interview tool', 'Your interview coach'],
            ['2. Setup', 'Upload your documents', 'Tell us about you and the role'],
            ['3. Questions', 'Here\'s what we\'ll test you on', 'Here\'s what we\'ll practice together'],
            ['4. Practice', 'Answer, then get help if stuck', 'Learn, then practice with confidence'],
            ['5. Completion', 'Here\'s your score', 'Here\'s how you\'ve grown'],
        ])

    # Journey comparison diagram
    add_image(
        doc,
        os.path.join(WIREFRAMES_PATH, 'journey-comparison.png'),
        'Figure 1: Test-First vs Teach-First user journey'
    )

    doc.add_page_break()


def phase_arrival(doc):
    add_heading(doc, 'Phase 1: Arrival')

    add_heading(doc, 'Current Experience', level=2)
    doc.add_paragraph(
        'User lands on the app. They see file upload fields and a "Generate Questions" action. '
        'The implicit message: this is a tool for testing yourself.'
    )

    add_image(
        doc,
        os.path.join(DIAGRAMS_PATH, 'state-01-empty.png'),
        'Figure 2: Current arrival state',
        width=5.0
    )

    add_heading(doc, 'Proposed Experience', level=2)
    doc.add_paragraph(
        'User lands on the app and immediately understands: this is a coach that helps you '
        'find and tell your story. The value proposition is clear before any action is taken.'
    )

    add_heading(doc, 'Recommendations', level=2)
    items = [
        'Lead with the value proposition: "Discover the stories in your experience"',
        'Show the journey: Setup → Learn → Practice → Grow',
        'Set expectations: "We\'ll use YOUR experience—nothing invented"',
        'Reduce visual complexity: progressive disclosure of options',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()


def phase_setup(doc):
    add_heading(doc, 'Phase 2: Setup')

    add_heading(doc, 'Current Experience', level=2)
    doc.add_paragraph(
        'User uploads resume file, uploads or pastes job description, clicks Generate. '
        'It feels transactional—providing inputs to a system.'
    )

    add_image(
        doc,
        os.path.join(DIAGRAMS_PATH, 'state-02-docs-uploaded.png'),
        'Figure 3: Current setup state',
        width=5.0
    )

    add_heading(doc, 'Proposed Experience', level=2)
    doc.add_paragraph(
        'User is guided through a conversation: "First, tell us about your background" '
        '(resume), then "Now, tell us about the role you\'re pursuing" (job). '
        'It feels like the start of a coaching session, not a form submission.'
    )

    add_heading(doc, 'Recommendations', level=2)
    items = [
        'Reframe file upload as "Share your background"',
        'Reframe job input as "Describe your target role"',
        'Add brief context: "We\'ll use this to personalize your coaching"',
        'Show progress: Step 1 of 3, Step 2 of 3...',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()


def phase_questions(doc):
    add_heading(doc, 'Phase 3: Questions')

    add_heading(doc, 'Current Experience', level=2)
    doc.add_paragraph(
        'Questions appear as a list. The user sees what they\'ll be "tested" on. '
        'Optional: they can add custom questions. Then they start the session.'
    )

    add_image(
        doc,
        os.path.join(DIAGRAMS_PATH, 'state-04-questions-timeout.png'),
        'Figure 4: Current questions state',
        width=5.0
    )

    add_heading(doc, 'Proposed Experience', level=2)
    doc.add_paragraph(
        'Questions are presented as "practice topics." For each question, the user can see: '
        'why this question matters for this role, and which parts of their resume are relevant. '
        'This teaches the connection BEFORE practice begins.'
    )

    add_heading(doc, 'Recommendations', level=2)
    items = [
        'Reframe: "Questions" → "Practice Topics" or "Coaching Areas"',
        'For each question, show the resume-to-question connection',
        'Let users preview: "What experience should I draw from?"',
        'Reduce anxiety: "You\'ll have guidance during practice"',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()


def phase_practice(doc):
    add_heading(doc, 'Phase 4: Practice')

    add_heading(doc, 'Current Experience', level=2)
    doc.add_paragraph(
        'Coach asks a question. User is expected to answer. If they freeze, they can seek help. '
        'The coaching moment comes AFTER the struggle.'
    )

    add_image(
        doc,
        os.path.join(DIAGRAMS_PATH, 'state-07-awaiting-answer.png'),
        'Figure 5: Current practice state - awaiting answer',
        width=5.0
    )

    # Current journey diagram
    add_image(
        doc,
        os.path.join(WIREFRAMES_PATH, 'journey-current.png'),
        'Figure 6: Current practice flow - teaching comes after freezing'
    )

    add_heading(doc, 'Proposed Experience', level=2)
    doc.add_paragraph(
        'Before each question, user sees how their resume answers it. '
        'They understand the narrative, then practice delivering it. '
        'The coaching moment comes BEFORE the performance.'
    )

    # Proposed journey diagram
    add_image(
        doc,
        os.path.join(WIREFRAMES_PATH, 'journey-proposed.png'),
        'Figure 7: Proposed practice flow - teaching comes before practicing'
    )

    add_heading(doc, 'Recommendations', level=2)
    items = [
        'Show relevant resume context before asking the question',
        'Frame the example as "Here\'s one way to answer using YOUR experience"',
        'Then invite: "Now try it in your own words"',
        'Feedback focuses on delivery, not content discovery',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()


def phase_completion(doc):
    add_heading(doc, 'Phase 5: Completion')

    add_heading(doc, 'Current Experience', level=2)
    doc.add_paragraph(
        'Session ends. User receives a score and can export a study guide. '
        'The framing is evaluative: here\'s how you did.'
    )

    add_image(
        doc,
        os.path.join(DIAGRAMS_PATH, 'state-08-score-display.png'),
        'Figure 8: Current completion state',
        width=5.0
    )

    add_heading(doc, 'Proposed Experience', level=2)
    doc.add_paragraph(
        'Session ends. User sees their growth: which narratives they discovered, '
        'which delivery improved, what to focus on next. The framing is developmental: '
        'here\'s how you\'ve grown and what to practice next.'
    )

    add_heading(doc, 'Recommendations', level=2)
    items = [
        'Reframe: "Score" → "Progress" or "Growth Summary"',
        'Highlight narratives discovered: "You found 4 strong stories"',
        'Provide specific next steps: "Practice these 2 topics again"',
        'Study guide becomes: "Your Interview Playbook"',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()


def strategic_options(doc):
    add_heading(doc, 'Implementation Options')

    doc.add_paragraph(
        'Three approaches to implementing the teach-first philosophy, '
        'particularly in the practice phase where the shift is most significant:'
    )

    add_heading(doc, 'Option A: Learning Mode', level=2)
    doc.add_paragraph(
        'Show the resume-grounded example BEFORE presenting the question. '
        'User learns the narrative, then practices delivering it.'
    )
    add_image(
        doc,
        os.path.join(WIREFRAMES_PATH, 'option-a-learning.png'),
        'Figure 9: Option A - Example first, then practice'
    )
    add_table(doc,
        ['Aspect', 'Detail'],
        [
            ['Best For', 'First-time users, high-anxiety users'],
            ['Trade-off', 'Confident users may want to try first'],
        ])

    add_heading(doc, 'Option B: Proactive Guidance', level=2)
    doc.add_paragraph(
        'Let users try first, but detect hesitation and offer guidance '
        'before they spiral into panic.'
    )
    add_image(
        doc,
        os.path.join(WIREFRAMES_PATH, 'option-b-proactive.png'),
        'Figure 10: Option B - Guidance offered proactively on hesitation'
    )
    add_table(doc,
        ['Aspect', 'Detail'],
        [
            ['Best For', 'Users who want to try first'],
            ['Trade-off', 'Requires tuning hesitation detection'],
        ])

    add_heading(doc, 'Option C: Always-Visible Context', level=2)
    doc.add_paragraph(
        'Show relevant resume context alongside the practice area at all times. '
        'User can glance at their experience while formulating answers.'
    )
    add_image(
        doc,
        os.path.join(WIREFRAMES_PATH, 'option-c-split.png'),
        'Figure 11: Option C - Resume context always visible'
    )
    add_table(doc,
        ['Aspect', 'Detail'],
        [
            ['Best For', 'Visual learners, reference-oriented users'],
            ['Trade-off', 'Reduced space for conversation flow'],
        ])

    doc.add_page_break()


def storyboard_section(doc):
    add_heading(doc, 'Full Storyboard: Redesigned Experience')

    doc.add_paragraph(
        'The following screens show the complete redesigned user journey using the teach-first '
        'coaching model with the proposed garnet color palette. Each screen demonstrates how '
        'the new positioning would feel to users.'
    )

    doc.add_paragraph()

    # Screen 1: Welcome
    add_heading(doc, 'Screen 1: Welcome', level=2)
    doc.add_paragraph(
        'The arrival experience positions PrepTalk as "Your Interview Coach" with warm, '
        'supportive messaging. The value proposition is clear: learn from examples, practice '
        'with your stories, build lasting confidence. No grades, no judgment, just growth.'
    )
    add_image(
        doc,
        os.path.join(STORYBOARD_PATH, '01-welcome.png'),
        'Figure 12: Welcome - "Your Interview Coach" positioning',
        width=5.5
    )

    # Screen 2: Setup
    add_heading(doc, 'Screen 2: Setup', level=2)
    doc.add_paragraph(
        'The setup phase guides users through sharing their background with coaching context. '
        '"I\'ll find the best stories from your experience" reframes file upload as the start '
        'of a coaching relationship, not a transactional form.'
    )
    add_image(
        doc,
        os.path.join(STORYBOARD_PATH, '02-setup.png'),
        'Figure 13: Setup - Guided upload with coaching context',
        width=5.5
    )

    # Screen 3: Questions Preview
    add_heading(doc, 'Screen 3: Questions Preview', level=2)
    doc.add_paragraph(
        'Questions are presented as "what we\'ll practice together" with learning mode visible '
        'and recommended. The coach avatar establishes the relationship before practice begins.'
    )
    add_image(
        doc,
        os.path.join(STORYBOARD_PATH, '03-questions.png'),
        'Figure 14: Questions Preview - Practice topics with learning mode',
        width=5.5
    )

    doc.add_page_break()

    # Screen 4: Learning Phase (KEY DIFFERENTIATOR)
    add_heading(doc, 'Screen 4: Learning Phase (Key Differentiator)', level=2)
    doc.add_paragraph(
        'THIS IS THE CORE TRANSFORMATION. Before each question, the user sees: (1) their '
        'relevant resume fact, (2) an example answer structure, and (3) why this approach works. '
        'They know which story to tell BEFORE being asked to tell it.'
    )
    add_image(
        doc,
        os.path.join(STORYBOARD_PATH, '04-learning.png'),
        'Figure 15: Learning Phase - Example BEFORE question prevents freeze',
        width=5.5
    )

    # Screen 5: Practice
    add_heading(doc, 'Screen 5: Practice with Context', level=2)
    doc.add_paragraph(
        'During practice, relevant resume cues remain visible on the right panel. '
        'The user is never alone—their experience is always within reach. '
        'Supportive status messages ("You\'ve got this") replace test anxiety.'
    )
    add_image(
        doc,
        os.path.join(STORYBOARD_PATH, '05-practice.png'),
        'Figure 16: Practice - Split panel with always-visible resume cues',
        width=5.5
    )

    doc.add_page_break()

    # Screen 6: Feedback
    add_heading(doc, 'Screen 6: Coaching Feedback', level=2)
    doc.add_paragraph(
        'Feedback leads with celebration ("Great job!") and what they did well. '
        'Constructive guidance follows. No scores—just growth-focused coaching tips. '
        'The STAR method suggestion is framed as "you naturally did most of this."'
    )
    add_image(
        doc,
        os.path.join(STORYBOARD_PATH, '06-feedback.png'),
        'Figure 17: Feedback - Celebrate first, then coach',
        width=5.5
    )

    # Screen 7: Completion
    add_heading(doc, 'Screen 7: Session Complete', level=2)
    doc.add_paragraph(
        'Completion is framed as growth: "Here\'s how you\'ve grown today." '
        'Key wins highlight specific achievements. Stats focus on stories ready to tell, '
        'not scores. The closing message reinforces confidence: "You\'re more prepared '
        'than you think. Trust your experience."'
    )
    add_image(
        doc,
        os.path.join(STORYBOARD_PATH, '07-complete.png'),
        'Figure 18: Completion - Growth summary, not score report',
        width=5.5
    )

    doc.add_page_break()


def recommendation(doc):
    add_heading(doc, 'Recommendation')

    p = doc.add_paragraph()
    p.add_run('Recommended: Start with Option A for new users, evolve based on data.').bold = True

    doc.add_paragraph()

    add_heading(doc, 'Rationale', level=2)
    items = [
        'Maximizes confidence for anxious users (our primary audience)',
        'Establishes PrepTalk\'s identity as coach, not judge',
        'Can graduate users to "try first" mode as they build confidence',
        'Lowest implementation complexity for highest strategic impact',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    add_heading(doc, 'Phased Rollout', level=2)
    add_table(doc,
        ['Phase', 'Focus', 'Deliverable'],
        [
            ['1', 'Arrival & Messaging', 'Updated copy, value proposition'],
            ['2', 'Practice Flow', 'Teach-first question experience'],
            ['3', 'Full Journey', 'Setup, questions, completion reframing'],
        ])

    add_heading(doc, 'Success Metrics', level=2)
    add_table(doc,
        ['Metric', 'Current', 'Target'],
        [
            ['Session completion rate', 'Baseline', 'Increase'],
            ['Time to first confident answer', 'Late in session', 'Earlier'],
            ['Return usage', 'Baseline', 'Increase'],
            ['User-reported confidence', 'Not measured', 'Track'],
        ])

    doc.add_page_break()


def decision_section(doc):
    add_heading(doc, 'Decision')

    doc.add_paragraph()

    question = doc.add_paragraph(
        'Should PrepTalk position itself as an interview simulator or a confidence builder?'
    )
    question.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in question.runs:
        run.italic = True
        run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    add_heading(doc, 'Options', level=2)
    options = [
        '☐ APPROVE: Proceed with teach-first repositioning as outlined',
        '☐ APPROVE WITH MODIFICATIONS: (specify below)',
        '☐ REQUEST MORE INFORMATION: (specify below)',
        '☐ DEFER: (specify timeline)',
    ]
    for opt in options:
        doc.add_paragraph(opt)

    doc.add_paragraph()
    doc.add_paragraph('Notes: _______________________________________________________')
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Signature: _______________________  Date: _______________')


def main():
    doc = Document()

    title_page(doc)
    executive_summary(doc)
    the_user(doc)
    design_principles(doc)
    journey_overview(doc)
    phase_arrival(doc)
    phase_setup(doc)
    phase_questions(doc)
    phase_practice(doc)
    phase_completion(doc)
    strategic_options(doc)
    storyboard_section(doc)
    recommendation(doc)
    decision_section(doc)

    doc.save(OUTPUT_PATH)
    print(f'Document saved: {OUTPUT_PATH}')


if __name__ == '__main__':
    main()
